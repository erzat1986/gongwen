#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公文排版脚本 - 根据国有资产AI盘活平台项目的正确格式进行修改
作者: Claude Code
版本: v1.1.0 (修复版本)
日期: 2024-05-10
"""

import argparse
import os
import shutil
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name, size_pt, bold=False, italic=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic

def is_level1_title(text):
    """判断是否为一级标题（一、）"""
    return re.match(r'^[一二三四五六七八九十]+、', text.strip())

def is_level2_title(text):
    """判断是否为二级标题（（一））"""
    return re.match(r'^（[一二三四五六七八九十]+）', text.strip())

def is_level3_title(text):
    """判断是否为三级标题（1.）"""
    return re.match(r'^\d+\.', text.strip())

def is_level4_title(text):
    """判断是否为四级标题（（1））"""
    return re.match(r'^（\d+）', text.strip())

def process_paragraph(para):
    """处理单个段落"""
    text = para.text.strip()

    # 清空现有格式
    para.clear()

    if not text:
        return

    # 判断是否为标题类型
    if any(keyword in text for keyword in ['项目建议书', '方案', '报告', '请示', '通知']):
        if len(text) < 30 and len(text.strip()) > 5:
            para.add_run(text)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = Pt(35)
            para.paragraph_format.first_line_indent = Cm(0)
            for run in para.runs:
                set_chinese_font(run, '方正小标宋_GBK', 22, bold=False)  # 大标题不加粗！！！
    elif is_level1_title(text):
        para.add_run(text)
        para.paragraph_format.first_line_indent = Cm(1.13)  # 修复缩进！！！
        para.paragraph_format.line_spacing = Pt(28)
        for run in para.runs:
            set_chinese_font(run, '黑体', 16, bold=True)
    elif is_level2_title(text):
        para.add_run(text)
        para.paragraph_format.first_line_indent = Cm(1.13)  # 修复缩进！！！
        para.paragraph_format.line_spacing = Pt(28)
        for run in para.runs:
            set_chinese_font(run, '楷体_GB2312', 16, bold=True)
    elif is_level3_title(text):
        para.add_run(text)
        para.paragraph_format.first_line_indent = Cm(1.13)  # 修复缩进！！！
        para.paragraph_format.line_spacing = Pt(28)
        for run in para.runs:
            set_chinese_font(run, '仿宋_GB2312', 16, bold=True)
    elif is_level4_title(text):
        para.add_run(text)
        para.paragraph_format.first_line_indent = Cm(1.13)  # 修复缩进！！！
        para.paragraph_format.line_spacing = Pt(28)
        for run in para.runs:
            set_chinese_font(run, '仿宋_GB2312', 16, bold=True)
    else:
        para.paragraph_format.first_line_indent = Cm(1.13)  # 修复缩进！！！
        para.paragraph_format.line_spacing = Pt(28)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 修复对齐！！！

        if text:
            run = para.add_run(text)
            set_chinese_font(run, '仿宋_GB2312', 16)

def format_document(input_path, output_path):
    """格式化Word文档"""
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"文件不存在: {input_path}")

    # 创建输出目录
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 备份原始文档
    backup_path = input_path.replace('.docx', '_原始备份.docx')
    if not os.path.exists(backup_path):
        shutil.copy2(input_path, backup_path)
        print(f"原始文档已备份至: {backup_path}")

    # 读取文档
    doc = Document(input_path)
    print(f"原始文档段落数: {len(doc.paragraphs)}")
    print(f"原始文档节数: {len(doc.sections)}")

    # 设置页面边距
    print("正在设置页面边距...")
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.7)
        section.right_margin = Cm(2.7)

    # 处理段落
    count = 0
    for para in doc.paragraphs:
        if para.text.strip():
            count += 1
        process_paragraph(para)

    print(f"已处理 {count} 个段落")

    # 保存文档
    doc.save(output_path)
    print(f"文档格式化完成！输出文件: {output_path}")

    return backup_path

def main():
    parser = argparse.ArgumentParser(description="公文排版脚本 - 根据国有资产AI盘活平台项目格式进行优化")
    parser.add_argument("input_file", help="输入的Word文档路径 (.docx)")
    parser.add_argument("-o", "--output", help="输出文件路径")

    args = parser.parse_args()

    input_path = args.input_file

    if not args.output:
        dir_name = os.path.dirname(input_path)
        file_name = os.path.basename(input_path)
        output_name = file_name.replace('.docx', '_公文排版版.docx')
        args.output = os.path.join(dir_name, output_name)

    try:
        backup_path = format_document(input_path, args.output)

        print("\n" + "="*50)
        print("公文排版任务完成！")
        print("="*50)
        print(f"原始文档: {input_path}")
        print(f"原始备份: {backup_path}")
        print(f"排版结果: {args.output}")
        print("="*50)

    except Exception as e:
        print(f"错误: {e}")
        import traceback
        print("详细错误信息:")
        print(traceback.format_exc())

if __name__ == "__main__":
    main()
