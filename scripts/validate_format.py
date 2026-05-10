#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
格式验证脚本 - 验证Word文档是否符合公文格式规范
"""

import os
import argparse
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def get_font_name(run):
    """获取字体名称"""
    try:
        return run.font.name
    except:
        return "未知"


def validate_document(doc_path):
    """验证文档格式"""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"文件不存在: {doc_path}")

    doc = Document(doc_path)
    issues = []

    print("正在验证公文格式...")
    print("=" * 60)

    # 1. 验证页面设置
    print("\n1. 页面设置验证:")
    sections = doc.sections
    if len(sections) > 0:
        section = sections[0]
        top_margin = section.top_margin.cm
        bottom_margin = section.bottom_margin.cm
        left_margin = section.left_margin.cm
        right_margin = section.right_margin.cm

        if not (3.6 <= top_margin <= 3.8):
            issues.append(f"上边距应为3.7cm，实际为{top_margin:.1f}cm")
        if not (3.4 <= bottom_margin <= 3.6):
            issues.append(f"下边距应为3.5cm，实际为{bottom_margin:.1f}cm")
        if not (2.6 <= left_margin <= 2.8):
            issues.append(f"左边距应为2.7cm，实际为{left_margin:.1f}cm")
        if not (2.6 <= right_margin <= 2.8):
            issues.append(f"右边距应为2.7cm，实际为{right_margin:.1f}cm")

        print(f"上边距: {top_margin:.1f}cm")
        print(f"下边距: {bottom_margin:.1f}cm")
        print(f"左边距: {left_margin:.1f}cm")
        print(f"右边距: {right_margin:.1f}cm")

    # 2. 验证段落格式
    print("\n2. 段落格式验证:")
    valid_titles = set()
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            has_text = False
            for run in para.runs:
                if run.text.strip():
                    has_text = True
                    break

            if has_text:
                indent = para.paragraph_format.first_line_indent.cm if para.paragraph_format.first_line_indent else 0
                line_spacing = para.paragraph_format.line_spacing.pt if hasattr(para.paragraph_format.line_spacing, 'pt') else "默认"

                # 检查标题
                text = para.text.strip()
                if len(text) < 30 and ('项目建议书' in text or '方案' in text or '报告' in text):
                    valid_titles.add(text)

                print(f"段落 {i+1:3d}: 长度={len(text):3d} 缩进={indent:.2f}cm 行距={line_spacing}")

    # 3. 验证字体使用
    print("\n3. 字体使用验证:")
    fonts = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip() and hasattr(run.font, 'name'):
                fonts.add(run.font.name)

    print(f"使用的字体: {', '.join(sorted(fonts))}")

    required_fonts = ['黑体', '楷体_GB2312', '仿宋_GB2312', 'Times New Roman']
    missing_fonts = []
    for font in required_fonts:
        if font not in fonts:
            missing_fonts.append(font)

    if missing_fonts:
        issues.append(f"缺少必需字体: {', '.join(missing_fonts)}")
        print(f"警告: 缺少必需字体: {', '.join(missing_fonts)}")

    # 4. 输出验证结果
    print("\n4. 验证结果:")
    if issues:
        print(f"发现 {len(issues)} 个格式问题:")
        for issue in issues:
            print(f"  - {issue}")
    else:
        print("所有格式符合公文规范！")

    print(f"\n文档共 {len(doc.paragraphs)} 个段落")
    print(f"标题识别: {len(valid_titles)} 个")

    return len(issues) == 0


def main():
    parser = argparse.ArgumentParser(description="公文格式验证脚本")
    parser.add_argument("doc_path", help="Word文档路径 (.docx)")

    args = parser.parse_args()

    try:
        valid = validate_document(args.doc_path)
        if valid:
            print("\n✅ 格式验证通过")
            return 0
        else:
            print("\n❌ 格式验证失败")
            return 1
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        print(traceback.format_exc())
        return 2


if __name__ == "__main__":
    exit(main())
